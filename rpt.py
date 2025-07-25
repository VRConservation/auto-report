import pandas as pd
from matplotlib import pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import argparse
from datetime import datetime
import os


class ReportGenerator:
    def __init__(self, excel_file='budget.xlsx', output_file=None):
        self.excel_file = excel_file

        # Create reports directory if it doesn't exist
        reports_dir = Path('reports')
        reports_dir.mkdir(exist_ok=True)

        # Simplified naming: project_report_yyyy-mm-dd.docx in reports folder
        if output_file is None:
            base_name = f'project_report_{datetime.now().strftime("%Y-%m-%d")}.docx'
            base_path = reports_dir / base_name
            self.output_file = str(self._get_unique_filename(base_path))
        else:
            # If custom filename provided, still put it in reports folder
            self.output_file = str(reports_dir / output_file)

        self.document = None
        self.data = None

    def _get_unique_filename(self, base_filepath):
        """Ensure unique filename by adding increment if file exists"""
        if not base_filepath.exists():
            return base_filepath
        
        # If file exists, add increment: project_report_2025-01-25_v2.docx
        name_stem = base_filepath.stem  # project_report_2025-01-25
        extension = base_filepath.suffix  # .docx
        parent_dir = base_filepath.parent
        
        counter = 2
        while True:
            new_filename = f"{name_stem}_v{counter}{extension}"
            new_filepath = parent_dir / new_filename
            if not new_filepath.exists():
                print(f"📝 File exists, creating new version: {new_filename}")
                return new_filepath
            counter += 1
            
            # Safety break to avoid infinite loop
            if counter > 100:
                # Use timestamp as fallback
                timestamp = datetime.now().strftime("%H%M%S")
                return parent_dir / f"{name_stem}_{timestamp}{extension}"
            
    def read_markdown_file(self, filename):
        """Read content from a markdown file if it exists"""
        try:
            file_path = Path(filename)
            if file_path.exists():
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read().strip()
            else:
                print(f"Info: {filename} not found, using default content")
                return None
        except Exception as e:
            print(f"Error reading {filename}: {e}")
            return None

    def add_markdown_content(self, filename, default_content=None):
        """Add content from markdown file to document with basic formatting"""
        content = self.read_markdown_file(filename)

        if not content and default_content:
            content = default_content
        elif not content:
            return

        # Split content into paragraphs
        paragraphs = content.split('\n\n')

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Handle bullet points (markdown style)
            if para.startswith('- ') or para.startswith('* '):
                # Split multiple bullet points
                bullets = [line.strip()[2:] for line in para.split(
                    '\n') if line.strip().startswith(('- ', '* '))]
                for bullet in bullets:
                    self.document.add_paragraph(bullet, style='List Bullet')
            else:
                # Regular paragraph
                self.document.add_paragraph(para)

    def load_data(self):
        """Load budget data from Excel file"""
        try:
            self.data = pd.read_excel(self.excel_file)
            print(f"✅ Successfully loaded data from {self.excel_file}")
            print(f"   Columns: {list(self.data.columns)}")
            print(f"   Rows: {len(self.data)}")
            return True
        except FileNotFoundError:
            print(f"❌ Error: Could not find {self.excel_file}")
            return False
        except Exception as e:
            print(f"❌ Error loading data: {e}")
            return False

    def create_document(self):
        """Initialize the Word document"""
        self.document = Document()
        # Add title with current date
        title = f'Budget Report - {datetime.now().strftime("%B %d, %Y")}'
        self.document.add_heading(title, 0)
        print("✅ Document initialized")

    def add_introduction(self):
        """Add introduction section"""
        self.document.add_heading('Executive Summary', level=1)

        default_intro = f"""This report provides a comprehensive overview of budget allocation and expenditure status as of {datetime.now().strftime("%B %d, %Y")}.

Key metrics include budget utilization rates, remaining fund allocation, and project-specific financial performance indicators."""

        self.add_markdown_content('introduction.md', default_intro)
        print("✅ Introduction section added")

    def add_budget_table(self):
        """Add budget data table"""
        if self.data is None:
            print("❌ No data available for table")
            return False

        self.document.add_heading('Budget Allocation Details', level=1)

        # Add summary paragraph
        total_budgeted = self.data['Budgeted'].sum(
        ) if 'Budgeted' in self.data.columns else 0
        total_remaining = self.data['Remaining'].sum(
        ) if 'Remaining' in self.data.columns else 0
        utilization_rate = ((total_budgeted - total_remaining) /
                            total_budgeted * 100) if total_budgeted > 0 else 0

        summary_text = f"Total Budget: ${total_budgeted:,.0f} | Utilization Rate: {utilization_rate:.1f}% | Remaining: ${total_remaining:,.0f}"
        self.document.add_paragraph(summary_text)

        # Create table
        table = self.document.add_table(rows=1, cols=len(self.data.columns))
        table.style = 'Table Grid'

        # Add header row with bold formatting
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(self.data.columns):
            hdr_cells[i].text = str(column_name)
            # Make header text bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for i in range(len(self.data)):
            row_cells = table.add_row().cells
            for j in range(len(self.data.columns)):
                cell_value = self.data.iloc[i, j]
                # Format numbers with commas if they're numeric
                if pd.api.types.is_numeric_dtype(type(cell_value)) and pd.notna(cell_value):
                    row_cells[j].text = f"{cell_value:,.0f}" if cell_value == int(
                        cell_value) else f"{cell_value:,.2f}"
                else:
                    row_cells[j].text = str(cell_value)

        print("✅ Budget table added")
        return True

    def add_key_points(self):
        """Add key points section"""
        self.document.add_heading('Key Findings', level=1)

        # Generate dynamic key points based on data
        key_points = []

        if self.data is not None and 'Budgeted' in self.data.columns and 'Remaining' in self.data.columns:
            # Calculate insights
            total_budgeted = self.data['Budgeted'].sum()
            total_remaining = self.data['Remaining'].sum()
            utilization_rate = ((total_budgeted - total_remaining) /
                                total_budgeted * 100) if total_budgeted > 0 else 0

            # Find highest and lowest utilization tasks
            self.data['Utilization%'] = (
                (self.data['Budgeted'] - self.data['Remaining']) / self.data['Budgeted'] * 100).round(1)

            if len(self.data) > 1:
                non_total_data = self.data[self.data['Task'] !='TOTALS'] if 'Task' in self.data.columns else self.data
                if len(non_total_data) > 0:
                    highest_util = non_total_data.loc[non_total_data['Utilization%'].idxmax(
                    )]
                    lowest_util = non_total_data.loc[non_total_data['Utilization%'].idxmin(
                    )]

                    key_points = [
                        f"Overall budget utilization stands at {utilization_rate:.1f}%",
                        f"Highest utilization: {highest_util['Task']} at {highest_util['Utilization%']:.1f}%",
                        f"Lowest utilization: {lowest_util['Task']} at {lowest_util['Utilization%']:.1f}%",
                        f"Total remaining funds: ${total_remaining:,.0f}"
                    ]

        # Use default if no dynamic points generated
        if not key_points:
            key_points = [
                "Budget tracking is current and accurate",
                "All expenditures are within approved parameters",
                "Financial controls are operating effectively",
                "Regular monitoring continues as scheduled"
            ]

        default_key_points = "\n".join([f"- {point}" for point in key_points])
        self.add_markdown_content('key_points.md', default_key_points)

        print("✅ Key points section added")

    def add_budget_chart(self):
        """Add budget visualization chart"""
        if self.data is None:
            print("❌ No data available for chart")
            return False

        self.document.add_heading('Budget Visualization', level=1)

        # Add chart description
        chart_desc = "The chart below provides a visual comparison of budgeted amounts versus remaining funds for each project component."
        self.add_markdown_content('chart_description.md', chart_desc)

        try:
            # Set matplotlib to non-interactive backend to prevent chart from showing
            plt.ioff()

            # Filter out TOTALS row for better visualization
            chart_data = self.data[self.data['Task'] !='TOTALS'] if 'Task' in self.data.columns else self.data

            if len(chart_data) == 0:
                chart_data = self.data

            # Create grouped bar chart
            ax = chart_data.plot(kind='bar',
                x='Task' if 'Task' in chart_data.columns else chart_data.index,
                y=['Budgeted', 'Remaining'],
                stacked=False,
            # Sea green and royal blue
                color=['#2E8B57', '#4169E1'],
                figsize=(12, 8))

            plt.title('Budget Status by Task', fontsize=16, fontweight='bold', pad=20)
            plt.xlabel('Project Tasks', fontsize=12)
            plt.ylabel('Amount ($)', fontsize=12)
            plt.xticks(rotation=45, ha='right')
            plt.legend(['Budgeted', 'Remaining'], loc='upper right')

            # Format y-axis with thousands separators
            ax.yaxis.set_major_formatter(
                plt.FuncFormatter(lambda x, p: f'{x:,.0f}'))

            # Add grid for better readability
            plt.grid(axis='y', alpha=0.3)

            plt.tight_layout()

            # Save chart
            chart_filename = 'budget_chart.png'
            plt.savefig(chart_filename, bbox_inches='tight',
                        dpi=300, facecolor='white')

            # Add to document
            self.document.add_picture(chart_filename, width=Inches(6.5))

            # Close the figure to prevent display and free memory
            plt.close()

            print("✅ Budget chart added")
            return True

        except Exception as e:
            print(f"❌ Error creating chart: {e}")
            return False

    def save_document(self):
        """Save the Word document"""
        try:
            self.document.save(self.output_file)
            print(f"✅ Report saved successfully: {self.output_file}")
            return True
        except Exception as e:
            print(f"❌ Error saving document: {e}")
            return False

    def generate_report(self):
        """Generate the complete report"""
        print("🚀 Starting report generation...")

        # Load data
        if not self.load_data():
            return False

        # Create document
        self.create_document()

        # Add sections
        self.add_introduction()
        self.add_budget_table()
        self.add_key_points()
        self.add_budget_chart()

        # Save document
        success = self.save_document()

        if success:
            print(f"🎉 Report generation completed successfully!")
            print(f"📄 Output file: {self.output_file}")

            # Print file size
            file_size = os.path.getsize(self.output_file) / 1024  # KB
            print(f"📊 File size: {file_size:.1f} KB")

        return success


def main():
    """Main function with command line argument support"""
    parser = argparse.ArgumentParser(
        description='Generate automated budget report')
    parser.add_argument('--input', '-i', default='budget.xlsx',
                        help='Input Excel file (default: budget.xlsx)')
    parser.add_argument('--output', '-o',
                        help='Output Word document filename')
    parser.add_argument('--verbose', '-v', action='store_true',
                        help='Enable verbose output')

    args = parser.parse_args()

    if args.verbose:
        print("🔧 Verbose mode enabled")
        print(f"📁 Input file: {args.input}")
        print(f"📄 Output file: {args.output or 'auto-generated'}")

    # Generate report
    generator = ReportGenerator(args.input, args.output)
    success = generator.generate_report()

    return 0 if success else 1


if __name__ == "__main__":
    exit(main())